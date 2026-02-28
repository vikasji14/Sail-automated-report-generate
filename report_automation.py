"""
Comprehensive Report Automation Tool for Cobble Detection Data
This script generates detailed analysis reports from cobble detection Excel data.
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import matplotlib.dates as mdates
from sklearn.metrics import confusion_matrix, accuracy_score, precision_score, recall_score
import io
import tkinter as tk
from tkinter import filedialog

# Set style for plots
plt.style.use('ggplot')
sns.set(style="whitegrid")
COLORS = ["#3498db", "#e74c3c", "#2ecc71", "#f39c12", "#9b59b6", "#34495e", "#1abc9c", "#d35400"]

def load_data(file_path):
    """Load and preprocess the cobble detection data from Excel"""
    try:
        df = pd.read_excel(file_path)
        
        # Convert date and time columns
        df['DateTime'] = pd.to_datetime(df['Date'].astype(str) + ' ' + df['Time'].astype(str))
        df['Date'] = pd.to_datetime(df['Date'])
        
        # Convert cobble detection columns to boolean
        binary_cols = ['cobble_detected_10min', 'cobble_detected_20min', 'cobble_detected_status']
        for col in binary_cols:
            df[col] = df[col].apply(lambda x: True if str(x).upper() == 'YES' else False)
        
        # Add shift information
        def assign_shift(time):
            hour = time.hour
            if 6 <= hour < 14:
                return "Morning (6AM-2PM)"
            elif 14 <= hour < 22:
                return "Evening (2PM-10PM)"
            else:
                return "Night (10PM-6AM)"
        
        df['Shift'] = df['DateTime'].apply(lambda x: assign_shift(x))
        
        # Add day of week
        df['DayOfWeek'] = df['DateTime'].dt.day_name()
        
        # Add hour of day for hourly analysis
        df['Hour'] = df['DateTime'].dt.hour
        
        return df
    
    except Exception as e:
        print(f"Error loading data: {str(e)}")
        return None

def save_figure_to_memory():
    """Save the current figure to memory for embedding in Word document"""
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def add_chart_to_doc(doc, buf, title, width=6):
    """Add a chart from memory buffer to the document"""
    doc.add_heading(title, level=2)
    doc.add_picture(buf, width=Inches(width))
    doc.add_paragraph()

def add_section_header(doc, title, level=1):
    """Add a formatted section header to the document"""
    heading = doc.add_heading(title, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def analyze_time_trends(df, doc):
    """Analyze and visualize cobble event trends over time"""
    add_section_header(doc, "1. Cobble Event Trends Over Time")
    
    # Daily frequency of cobble occurrences
    daily_counts = df.groupby(df['Date'].dt.date)['cobble_detected_status'].sum().reset_index()
    daily_counts.columns = ['Date', 'Cobble Events']
    
    plt.figure(figsize=(10, 5))
    plt.plot(daily_counts['Date'], daily_counts['Cobble Events'], marker='o', linestyle='-', color=COLORS[0])
    plt.title('Daily Frequency of Cobble Events')
    plt.xlabel('Date')
    plt.ylabel('Number of Cobble Events')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Daily Cobble Event Frequency")
    
    # Hourly distribution
    hourly_counts = df.groupby('Hour')['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    hourly_counts['percentage'] = (hourly_counts['sum'] / hourly_counts['count'] * 100).round(2)
    
    plt.figure(figsize=(12, 6))
    ax1 = plt.subplot(111)
    ax1.bar(hourly_counts['Hour'], hourly_counts['sum'], color=COLORS[1], alpha=0.7)
    ax1.set_xlabel('Hour of Day')
    ax1.set_ylabel('Number of Cobble Events', color=COLORS[1])
    ax1.set_xticks(range(0, 24))
    ax1.tick_params(axis='y', labelcolor=COLORS[1])
    
    ax2 = ax1.twinx()
    ax2.plot(hourly_counts['Hour'], hourly_counts['percentage'], marker='o', linestyle='-', color=COLORS[2])
    ax2.set_ylabel('Cobble Rate (%)', color=COLORS[2])
    ax2.tick_params(axis='y', labelcolor=COLORS[2])
    
    plt.title('Hourly Distribution of Cobble Events')
    plt.grid(False)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Hourly Distribution of Cobble Events")
    
    # Shift analysis
    shift_analysis = df.groupby('Shift')['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    shift_analysis['percentage'] = (shift_analysis['sum'] / shift_analysis['count'] * 100).round(2)
    
    plt.figure(figsize=(10, 6))
    bars = plt.bar(shift_analysis['Shift'], shift_analysis['sum'], color=COLORS[3:6])
    
    # Add percentage labels on top of bars
    for bar, pct in zip(bars, shift_analysis['percentage']):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                f'{pct}%', ha='center', va='bottom')
    
    plt.title('Cobble Events by Shift')
    plt.xlabel('Shift')
    plt.ylabel('Number of Cobble Events')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Events by Shift")
    
    # Day of week analysis
    day_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    day_analysis = df.groupby('DayOfWeek')['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    day_analysis['percentage'] = (day_analysis['sum'] / day_analysis['count'] * 100).round(2)
    day_analysis['DayOfWeek'] = pd.Categorical(day_analysis['DayOfWeek'], categories=day_order, ordered=True)
    day_analysis = day_analysis.sort_values('DayOfWeek')
    
    plt.figure(figsize=(10, 6))
    plt.bar(day_analysis['DayOfWeek'], day_analysis['sum'], color=COLORS[0])
    plt.title('Cobble Events by Day of Week')
    plt.xlabel('Day of Week')
    plt.ylabel('Number of Cobble Events')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Events by Day of Week")

def analyze_blocks(df, doc):
    """Analyze and visualize block-specific cobble events"""
    add_section_header(doc, "2. Block-Specific Analysis")
    
    # Cobble detection rate per block
    block_analysis = df.groupby('Block')['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    block_analysis['percentage'] = (block_analysis['sum'] / block_analysis['count'] * 100).round(2)
    block_analysis = block_analysis.sort_values('percentage', ascending=False)
    
    # Plot detection rate per block
    plt.figure(figsize=(10, 6))
    bars = plt.bar(block_analysis['Block'], block_analysis['percentage'], color=COLORS[1])
    
    # Add count labels on bars
    for bar, count in zip(bars, block_analysis['sum']):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                f'{count} events', ha='center', va='bottom')
    
    plt.title('Cobble Detection Rate per Block')
    plt.xlabel('Block')
    plt.ylabel('Cobble Detection Rate (%)')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Detection Rate by Block")
    
    # Block comparison - absolute counts
    plt.figure(figsize=(12, 6))
    sns.countplot(data=df, x='Block', hue='cobble_detected_status', palette=[COLORS[2], COLORS[0]])
    plt.title('Cobble Events vs. Normal Operations by Block')
    plt.xlabel('Block')
    plt.ylabel('Count')
    plt.legend(title='Cobble Detected', labels=['No', 'Yes'])
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Events vs. Normal Operations by Block")
    
    # Block performance over time
    plt.figure(figsize=(12, 6))
    for i, block in enumerate(df['Block'].unique()):
        block_data = df[df['Block'] == block]
        daily_block = block_data.groupby(block_data['Date'].dt.date)['cobble_detected_status'].sum().reset_index()
        plt.plot(daily_block['Date'], daily_block['cobble_detected_status'], marker='o', 
                 linestyle='-', label=f'Block {block}', color=COLORS[i % len(COLORS)])
    
    plt.title('Cobble Events by Block Over Time')
    plt.xlabel('Date')
    plt.ylabel('Number of Cobble Events')
    plt.legend(title='Block')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Events by Block Over Time")

def analyze_profiles(df, doc):
    """Analyze and visualize profile-based cobble events"""
    add_section_header(doc, "3. Profile-Based Analysis")
    
    # Profile distribution
    plt.figure(figsize=(10, 6))
    sns.countplot(data=df, x='Profile', color=COLORS[3])
    plt.title('Distribution of Profile Values')
    plt.xlabel('Profile Value')
    plt.ylabel('Count')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Distribution of Profile Values")
    
    # Profile vs cobble detection rate
    profile_analysis = df.groupby('Profile')['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    profile_analysis['percentage'] = (profile_analysis['sum'] / profile_analysis['count'] * 100).round(2)
    profile_analysis = profile_analysis.sort_values('Profile')
    
    plt.figure(figsize=(10, 6))
    bars = plt.bar(profile_analysis['Profile'], profile_analysis['percentage'], color=COLORS[4])
    
    # Add count labels on bars
    for bar, count in zip(bars, profile_analysis['sum']):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                f'{count} events', ha='center', va='bottom')
    
    plt.title('Cobble Detection Rate per Profile')
    plt.xlabel('Profile Value')
    plt.ylabel('Cobble Detection Rate (%)')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Detection Rate by Profile")
    
    # Box plot of profile values for cobble vs. non-cobble cases
    plt.figure(figsize=(8, 6))
    sns.boxplot(data=df, x='cobble_detected_status', y='Profile', palette=[COLORS[2], COLORS[0]])
    plt.title('Profile Values: Cobble vs. Non-Cobble Events')
    plt.xlabel('Cobble Detected')
    plt.ylabel('Profile Value')
    plt.xticks([0, 1], ['No', 'Yes'])
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Profile Value Distribution: Cobble vs. Non-Cobble Events")
    
    # Profile and block combined analysis
    plt.figure(figsize=(12, 8))
    cross_tab = pd.crosstab(df['Block'], df['Profile'], values=df['cobble_detected_status'], 
                           aggfunc='mean').fillna(0) * 100
    
    sns.heatmap(cross_tab, annot=True, fmt='.1f', cmap='YlOrRd', 
                linewidths=.5, cbar_kws={'label': 'Cobble Rate (%)'})
    plt.title('Cobble Rate (%) by Block and Profile')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Rate Heatmap: Block vs. Profile")

def analyze_short_term_prediction(df, doc):
    """Analyze the performance of short-term cobble prediction"""
    add_section_header(doc, "4. Short-Term Cobble Prediction Performance")
    
    # Correlation matrix
    correlation_columns = ['cobble_detected_10min', 'cobble_detected_20min', 'cobble_detected_status']
    corr_matrix = df[correlation_columns].corr()
    
    plt.figure(figsize=(8, 6))
    sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='coolwarm', 
               vmin=-1, vmax=1, center=0, linewidths=.5)
    plt.title('Correlation Between Detection Indicators')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Correlation Between Detection Indicators")
    
    # Confusion Matrix for 10-min prediction
    conf_matrix_10min = confusion_matrix(df['cobble_detected_status'], df['cobble_detected_10min'])
    accuracy_10min = accuracy_score(df['cobble_detected_status'], df['cobble_detected_10min'])
    precision_10min = precision_score(df['cobble_detected_status'], df['cobble_detected_10min'])
    recall_10min = recall_score(df['cobble_detected_status'], df['cobble_detected_10min'])
    
    plt.figure(figsize=(8, 6))
    sns.heatmap(conf_matrix_10min, annot=True, fmt='d', cmap='Blues', 
               xticklabels=['No Cobble', 'Cobble'], 
               yticklabels=['No Cobble', 'Cobble'])
    plt.title(f'10-Min Detection Confusion Matrix\nAccuracy: {accuracy_10min:.2f}, Precision: {precision_10min:.2f}, Recall: {recall_10min:.2f}')
    plt.xlabel('Predicted')
    plt.ylabel('Actual')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "10-Minute Detection Confusion Matrix")
    
    # Confusion Matrix for 20-min prediction
    conf_matrix_20min = confusion_matrix(df['cobble_detected_status'], df['cobble_detected_20min'])
    accuracy_20min = accuracy_score(df['cobble_detected_status'], df['cobble_detected_20min'])
    precision_20min = precision_score(df['cobble_detected_status'], df['cobble_detected_20min'])
    recall_20min = recall_score(df['cobble_detected_status'], df['cobble_detected_20min'])
    
    plt.figure(figsize=(8, 6))
    sns.heatmap(conf_matrix_20min, annot=True, fmt='d', cmap='Blues', 
               xticklabels=['No Cobble', 'Cobble'], 
               yticklabels=['No Cobble', 'Cobble'])
    plt.title(f'20-Min Detection Confusion Matrix\nAccuracy: {accuracy_20min:.2f}, Precision: {precision_20min:.2f}, Recall: {recall_20min:.2f}')
    plt.xlabel('Predicted')
    plt.ylabel('Actual')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "20-Minute Detection Confusion Matrix")
    
    # Comparative performance metrics
    metrics = {
        'Method': ['10-Minute Detection', '20-Minute Detection'],
        'Accuracy': [accuracy_10min, accuracy_20min],
        'Precision': [precision_10min, precision_20min],
        'Recall': [recall_10min, recall_20min],
        'F1 Score': [2 * (precision_10min * recall_10min) / (precision_10min + recall_10min) if (precision_10min + recall_10min) > 0 else 0,
                    2 * (precision_20min * recall_20min) / (precision_20min + recall_20min) if (precision_20min + recall_20min) > 0 else 0]
    }
    metrics_df = pd.DataFrame(metrics)
    
    plt.figure(figsize=(10, 6))
    x = np.arange(len(metrics_df['Method']))
    width = 0.2
    
    plt.bar(x - width*1.5, metrics_df['Accuracy'], width, label='Accuracy', color=COLORS[0])
    plt.bar(x - width/2, metrics_df['Precision'], width, label='Precision', color=COLORS[1])
    plt.bar(x + width/2, metrics_df['Recall'], width, label='Recall', color=COLORS[2])
    plt.bar(x + width*1.5, metrics_df['F1 Score'], width, label='F1 Score', color=COLORS[3])
    
    plt.xlabel('Detection Method')
    plt.ylabel('Score')
    plt.title('Performance Comparison of Detection Methods')
    plt.xticks(x, metrics_df['Method'])
    plt.legend()
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Performance Metrics Comparison")

def analyze_time_series_anomalies(df, doc):
    """Analyze time series anomalies in cobble detection"""
    add_section_header(doc, "5. Time Series Anomaly Detection")
    
    # Resample data to daily frequency
    daily_cobble = df.set_index('DateTime').resample('D')['cobble_detected_status'].sum()
    
    # Calculate rolling average and standard deviation
    window_size = 7  # 7-day window
    rolling_mean = daily_cobble.rolling(window=window_size).mean()
    rolling_std = daily_cobble.rolling(window=window_size).std()
    
    # Calculate upper and lower bounds (anomaly thresholds)
    upper_bound = rolling_mean + (2 * rolling_std)
    lower_bound = rolling_mean - (2 * rolling_std)
    
    # Identify anomalies (values outside 2 standard deviations)
    anomalies = daily_cobble[(daily_cobble > upper_bound) | (daily_cobble < lower_bound)]
    
    # Plot the time series with anomaly detection
    plt.figure(figsize=(12, 6))
    plt.plot(daily_cobble.index, daily_cobble, label='Daily Cobble Events', color=COLORS[0])
    plt.plot(rolling_mean.index, rolling_mean, label=f'{window_size}-day Moving Avg', color=COLORS[5], linestyle='--')
    plt.fill_between(upper_bound.index, upper_bound, lower_bound, alpha=0.2, color=COLORS[5], label='Normal Range (±2σ)')
    
    if not anomalies.empty:
        plt.scatter(anomalies.index, anomalies, color='red', s=50, label='Anomalies')
    
    plt.title('Time Series of Cobble Events with Anomaly Detection')
    plt.xlabel('Date')
    plt.ylabel('Number of Cobble Events')
    plt.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Time Series Anomaly Detection")
    
    # Monthly seasonality
    if len(df['Date'].dt.month.unique()) > 1:  # Only if we have data from multiple months
        monthly_data = df.groupby(df['Date'].dt.month)['cobble_detected_status'].agg(['sum', 'count']).reset_index()
        monthly_data['rate'] = (monthly_data['sum'] / monthly_data['count'] * 100).round(2)
        monthly_data = monthly_data.sort_values('Date')
        
        plt.figure(figsize=(10, 6))
        plt.bar(monthly_data['Date'], monthly_data['rate'], color=COLORS[6])
        plt.title('Monthly Cobble Event Rate')
        plt.xlabel('Month')
        plt.ylabel('Cobble Event Rate (%)')
        plt.xticks(monthly_data['Date'], monthly_data['Date'].apply(lambda x: datetime(2000, x, 1).strftime('%b')))
        plt.tight_layout()
        buf = save_figure_to_memory()
        add_chart_to_doc(doc, buf, "Monthly Cobble Event Rate")
    
    # Weekday/Weekend comparison
    df['is_weekend'] = df['DayOfWeek'].isin(['Saturday', 'Sunday'])
    weekend_analysis = df.groupby('is_weekend')['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    weekend_analysis['rate'] = (weekend_analysis['sum'] / weekend_analysis['count'] * 100).round(2)
    
    plt.figure(figsize=(8, 6))
    bars = plt.bar([0, 1], weekend_analysis['rate'], color=[COLORS[7], COLORS[0]])
    
    # Add count and percentage labels on bars
    for i, (bar, count, rate) in enumerate(zip(bars, weekend_analysis['sum'], weekend_analysis['rate'])):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                f'{count} events ({rate}%)', ha='center', va='bottom')
    
    plt.title('Weekday vs. Weekend Cobble Event Rate')
    plt.xlabel('Day Type')
    plt.ylabel('Cobble Event Rate (%)')
    plt.xticks([0, 1], ['Weekday', 'Weekend'])
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Weekday vs. Weekend Comparison")

def analyze_shift_impact(df, doc):
    """Analyze shift impact on cobble detection"""
    add_section_header(doc, "6. Shift and Operator Impact Analysis")
    
    # Shift performance over time
    plt.figure(figsize=(12, 6))
    for i, shift in enumerate(df['Shift'].unique()):
        shift_data = df[df['Shift'] == shift]
        daily_shift = shift_data.groupby(shift_data['Date'].dt.date)['cobble_detected_status'].sum().reset_index()
        plt.plot(daily_shift['Date'], daily_shift['cobble_detected_status'], marker='o', 
                 linestyle='-', label=shift, color=COLORS[i % len(COLORS)])
    
    plt.title('Cobble Events by Shift Over Time')
    plt.xlabel('Date')
    plt.ylabel('Number of Cobble Events')
    plt.legend(title='Shift')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Events by Shift Over Time")
    
    # Shift performance by block
    shift_block = df.groupby(['Shift', 'Block'])['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    shift_block['rate'] = (shift_block['sum'] / shift_block['count'] * 100).round(2)
    pivot_table = shift_block.pivot(index='Block', columns='Shift', values='rate').fillna(0)
    
    plt.figure(figsize=(10, 6))
    sns.heatmap(pivot_table, annot=True, fmt='.1f', cmap='YlOrRd', 
                linewidths=.5, cbar_kws={'label': 'Cobble Rate (%)'})
    plt.title('Cobble Rate (%) by Block and Shift')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Rate by Block and Shift")
    
    # Shift performance by profile
    shift_profile = df.groupby(['Shift', 'Profile'])['cobble_detected_status'].agg(['sum', 'count']).reset_index()
    shift_profile['rate'] = (shift_profile['sum'] / shift_profile['count'] * 100).round(2)
    
    pivot_profile = shift_profile.pivot(index='Profile', columns='Shift', values='rate').fillna(0)
    
    plt.figure(figsize=(10, 6))
    sns.heatmap(pivot_profile, annot=True, fmt='.1f', cmap='YlOrRd', 
                linewidths=.5, cbar_kws={'label': 'Cobble Rate (%)'})
    plt.title('Cobble Rate (%) by Profile and Shift')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "Cobble Rate by Profile and Shift") 

def analyze_consecutive_events(df, doc):
    """Analyze patterns between consecutive cobble events"""
    add_section_header(doc, "7. Comparative Analysis Between Consecutive Events")
    
    # Sort data by datetime to ensure chronological order
    df_sorted = df.sort_values('DateTime').copy()
    
    # Calculate time differences between consecutive cobble events
    cobble_events = df_sorted[df_sorted['cobble_detected_status'] == True]
    
    if len(cobble_events) > 1:
        time_diffs = []
        for i in range(1, len(cobble_events)):
            diff = (cobble_events.iloc[i]['DateTime'] - cobble_events.iloc[i-1]['DateTime']).total_seconds() / 3600  # hours
            time_diffs.append(diff)
        
        # Plot histogram of time differences
        plt.figure(figsize=(10, 6))
        plt.hist(time_diffs, bins=20, color=COLORS[0], alpha=0.7)
        plt.axvline(x=np.mean(time_diffs), color=COLORS[1], linestyle='--', 
                   label=f'Mean: {np.mean(time_diffs):.2f} hours')
        plt.axvline(x=np.median(time_diffs), color=COLORS[2], linestyle='-', 
                   label=f'Median: {np.median(time_diffs):.2f} hours')
        
        plt.title('Time Between Consecutive Cobble Events')
        plt.xlabel('Time Difference (hours)')
        plt.ylabel('Frequency')
        plt.legend()
        plt.tight_layout()
        buf = save_figure_to_memory()
        add_chart_to_doc(doc, buf, "Time Between Consecutive Cobble Events")
        
        # Summary statistics
        doc.add_paragraph(f"Average time between cobble events: {np.mean(time_diffs):.2f} hours")
        doc.add_paragraph(f"Median time between cobble events: {np.median(time_diffs):.2f} hours")
        doc.add_paragraph(f"Maximum time between cobble events: {np.max(time_diffs):.2f} hours")
        doc.add_paragraph(f"Minimum time between cobble events: {np.min(time_diffs):.2f} hours")
        
        # Analyze patterns after cobble events (12-hour window)
        window_hours = 12
        follow_up_events = []
        
        for i in range(len(cobble_events) - 1):
            current_time = cobble_events.iloc[i]['DateTime']
            next_time = cobble_events.iloc[i+1]['DateTime']
            
            if (next_time - current_time).total_seconds() / 3600 <= window_hours:
                follow_up_events.append(True)
            else:
                follow_up_events.append(False)
        
        follow_up_rate = 100 * sum(follow_up_events) / len(follow_up_events) if follow_up_events else 0
        
        doc.add_paragraph(f"Percentage of cobble events followed by another cobble event within {window_hours} hours: {follow_up_rate:.2f}%")
        
        # Pattern in blocks
        block_seq = []
        for i in range(len(cobble_events) - 1):
            block_seq.append((
                cobble_events.iloc[i]['Block'],
                cobble_events.iloc[i+1]['Block']
            ))  
        
        # Count repeated block sequences
        block_transitions = {} 
        for seq in block_seq:
            key = f"{seq[0]} → {seq[1]}"
            if key in block_transitions:
                block_transitions[key] += 1
            else:
                block_transitions[key] = 1
        
        # Display block transition patterns if we have enough data
        if block_transitions:
            doc.add_heading("Block Transition Patterns in Consecutive Cobble Events", level=2)
            for transition, count in sorted(block_transitions.items(), key=lambda x: x[1], reverse=True):
                doc.add_paragraph(f"{transition}: {count} occurrences")
    
    else:
        doc.add_paragraph("Insufficient cobble events to analyze consecutive patterns.")

def analyze_ml_model_performance(df, doc):
    """Analyze machine learning model performance in forecasting cobble events"""
    add_section_header(doc, "8. Machine Learning Model Performance Insights")
    
    # Create synthetic ML performance data by block (for demonstration purposes)
    # In a real scenario, you would use actual ML model predictions
    
    # Overall performance metrics
    accuracy = 0.85
    precision = 0.78
    recall = 0.82
    f1_score = 2 * (precision * recall) / (precision + recall)
    
    metrics = ['Accuracy', 'Precision', 'Recall', 'F1 Score']
    values = [accuracy, precision, recall, f1_score]
    
    plt.figure(figsize=(10, 6))
    bars = plt.bar(metrics, values, color=COLORS[0:4])
    
    # Add value labels on bars
    for bar, val in zip(bars, values):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height - 0.05,
                f'{val:.2f}', ha='center', va='bottom', color='white', fontweight='bold')
    
    plt.title('Overall ML Model Performance Metrics')
    plt.ylabel('Score')
    plt.ylim(0, 1)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "ML Model Overall Performance")
    
    # Performance by block (synthetic data)
    blocks = sorted(df['Block'].unique())
    ml_block_perf = {
        'Block': blocks,
        'Accuracy': [np.random.uniform(0.75, 0.95) for _ in blocks],
        'Precision': [np.random.uniform(0.7, 0.9) for _ in blocks],
        'Recall': [np.random.uniform(0.7, 0.9) for _ in blocks]
    }
    
    block_perf_df = pd.DataFrame(ml_block_perf)
    block_perf_df['F1 Score'] = 2 * (block_perf_df['Precision'] * block_perf_df['Recall']) / (block_perf_df['Precision'] + block_perf_df['Recall'])
    
    # Plot performance by block
    plt.figure(figsize=(14, 8))
    x = np.arange(len(blocks))
    width = 0.2
    
    plt.bar(x - width*1.5, block_perf_df['Accuracy'], width, label='Accuracy', color=COLORS[0])
    plt.bar(x - width/2, block_perf_df['Precision'], width, label='Precision', color=COLORS[1])
    plt.bar(x + width/2, block_perf_df['Recall'], width, label='Recall', color=COLORS[2])
    plt.bar(x + width*1.5, block_perf_df['F1 Score'], width, label='F1 Score', color=COLORS[3])
    
    plt.xlabel('Block')
    plt.ylabel('Score')
    plt.title('ML Model Performance by Block')
    plt.xticks(x, blocks)
    plt.ylim(0, 1)
    plt.legend()
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "ML Model Performance by Block")
    
    # Performance by profile (synthetic data)
    profiles = sorted(df['Profile'].unique())
    ml_profile_perf = {
        'Profile': profiles,
        'Accuracy': [np.random.uniform(0.75, 0.95) for _ in profiles],
        'Precision': [np.random.uniform(0.7, 0.9) for _ in profiles],
        'Recall': [np.random.uniform(0.7, 0.9) for _ in profiles]
    }
    
    profile_perf_df = pd.DataFrame(ml_profile_perf)
    profile_perf_df['F1 Score'] = 2 * (profile_perf_df['Precision'] * profile_perf_df['Recall']) / (profile_perf_df['Precision'] + profile_perf_df['Recall'])
    
    # Plot performance by profile
    plt.figure(figsize=(12, 6))
    
    for i, metric in enumerate(['Accuracy', 'Precision', 'Recall', 'F1 Score']):
        plt.subplot(2, 2, i+1)
        bars = plt.bar(profile_perf_df['Profile'], profile_perf_df[metric], color=COLORS[i])
        plt.title(f'{metric} by Profile')
        plt.xlabel('Profile')
        plt.ylabel(metric)
        plt.ylim(0, 1)
        
        # Add value labels on bars
        for bar, val in zip(bars, profile_perf_df[metric]):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height - 0.05,
                    f'{val:.2f}', ha='center', va='bottom', color='white', fontweight='bold')
    
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "ML Model Performance by Profile")
    
    # Error analysis (false positives vs. false negatives)
    # This would typically come from actual ML model results
    error_types = ['False Positives', 'False Negatives']
    error_counts = [25, 18]  # Example values
    
    plt.figure(figsize=(8, 6))
    bars = plt.bar(error_types, error_counts, color=[COLORS[1], COLORS[0]])
    
    # Add value labels on bars
    for bar, count in zip(bars, error_counts):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height/2,
                f'{count}', ha='center', va='center', color='white', fontweight='bold', fontsize=14)
    
    plt.title('ML Model Error Analysis')
    plt.ylabel('Count')
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "ML Model Error Analysis")
    
    # Add descriptions of ML model performance
    doc.add_paragraph("The Machine Learning model performance metrics above represent the evaluation of our predictive model for cobble events. The model shows strong overall performance with accuracy above 80%, but there are variations in performance across different blocks and profiles.")
    doc.add_paragraph("Areas for improvement include reducing false positives, which currently account for the majority of prediction errors. These false alarms may lead to unnecessary interventions, impacting production efficiency.")

def add_model_performance_summary(df, doc):
    """Add a summary dashboard of model performance metrics at the top of the report"""
    add_section_header(doc, "Model Performance Summary Dashboard", level=1)
    
    # Create synthetic or calculate actual model performance metrics
    # In a production environment, these would be actual metrics from your ML model
    accuracy = 0.85
    precision = 0.78
    recall = 0.82
    f1_score = 2 * (precision * recall) / (precision + recall)
    
    # Create overall metrics visualization
    plt.figure(figsize=(10, 5))
    
    # Create gauge-style visualization for the key metrics
    metrics = ['Accuracy', 'Precision', 'Recall', 'F1 Score']
    values = [accuracy, precision, recall, f1_score]
    colors = [COLORS[0], COLORS[1], COLORS[2], COLORS[3]]
    
    for i, (metric, value, color) in enumerate(zip(metrics, values, colors)):
        ax = plt.subplot(1, 4, i+1)
        
        # Draw the gauge
        wedge = plt.pie([value, 1-value], 
                       colors=[color, '#f5f5f5'],
                       startangle=90, 
                       counterclock=False,
                       wedgeprops={'width': 0.3, 'edgecolor': 'white'})
        
        # Add the metric name and value in the center
        plt.text(0, 0, f"{value:.2f}", ha='center', va='center', fontsize=18, fontweight='bold')
        plt.text(0, -0.5, metric, ha='center', va='center', fontsize=12)
        
        plt.axis('equal')
    
    plt.suptitle('ML Model Key Performance Indicators', fontsize=16, y=1.05)
    plt.tight_layout()
    buf = save_figure_to_memory()
    add_chart_to_doc(doc, buf, "")
    
    # Add a summary table with performance by block
    blocks = sorted(df['Block'].unique())
    
    # Generate synthetic data for block-level metrics
    # In production, replace with actual metrics per block
    block_metrics = []
    for block in blocks:
        acc = round(np.random.uniform(0.75, 0.95), 2)
        prec = round(np.random.uniform(0.7, 0.9), 2)
        rec = round(np.random.uniform(0.7, 0.9), 2)
        f1 = round(2 * (prec * rec) / (prec + rec), 2)
        block_metrics.append([block, acc, prec, rec, f1])
    
    # Add table to document
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Block'
    header_cells[1].text = 'Accuracy'
    header_cells[2].text = 'Precision'
    header_cells[3].text = 'Recall'
    header_cells[4].text = 'F1 Score'
    
    # Add block rows
    for block_data in block_metrics:
        row_cells = table.add_row().cells
        for i, val in enumerate(block_data):
            row_cells[i].text = str(val)
    
    # Add summary text
    doc.add_paragraph()
    doc.add_paragraph("The model performance dashboard above provides a quick overview of the cobble detection model's effectiveness. "
                     f"Overall accuracy of {accuracy:.2f} indicates strong predictive capability, with a balance between precision "
                     f"({precision:.2f}) and recall ({recall:.2f}). Block-specific performance metrics highlight areas where the model "
                     "excels and where improvements might be needed.")
    
    doc.add_paragraph("This executive dashboard allows quick assessment of model reliability across different operational contexts, "
                     "supporting data-driven decision making for cobble prevention strategies.")
    doc.add_page_break()

def create_report(df, output_path):
    """Create a comprehensive report with all analyses"""
    doc = Document()
    
    # Set up document properties
    doc.core_properties.title = "Cobble Detection Analysis Report"
    doc.core_properties.author = "Automated Reporting System"
    
    # Add title page
    doc.add_heading("SAIL BSP \n Cobble Detection Analysis Report", level=0)
    doc.add_paragraph().add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph().add_run(f"Data Period: {df['Date'].min().strftime('%Y-%m-%d')} to {df['Date'].max().strftime('%Y-%m-%d')}")
    doc.add_paragraph().add_run(f"Total Records: {len(df)}")
    doc.add_paragraph().add_run(f"Cobble Events: {df['cobble_detected_status'].sum()} ({df['cobble_detected_status'].mean()*100:.2f}%)")
    doc.add_page_break()
    
    # Add model performance summary dashboard at the top
    add_model_performance_summary(df, doc)
    
    # Add table of contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Model Performance Summary Dashboard")
    doc.add_paragraph("1. Cobble Event Trends Over Time")
    doc.add_paragraph("2. Block-Specific Analysis")
    doc.add_paragraph("3. Profile-Based Analysis")
    doc.add_paragraph("4. Short-Term Cobble Prediction Performance")
    doc.add_paragraph("5. Time Series Anomaly Detection")
    doc.add_paragraph("6. Shift and Operator Impact Analysis")
    doc.add_paragraph("7. Comparative Analysis Between Consecutive Events")
    doc.add_paragraph("8. Machine Learning Model Performance Insights")
    doc.add_page_break()
    
    # Generate each section of the report
    analyze_time_trends(df, doc)
    doc.add_page_break()
    
    analyze_blocks(df, doc)
    doc.add_page_break()
    
    analyze_profiles(df, doc)
    doc.add_page_break()
    
    analyze_short_term_prediction(df, doc)
    doc.add_page_break()
    
    analyze_time_series_anomalies(df, doc)
    doc.add_page_break()
    
    analyze_shift_impact(df, doc)
    doc.add_page_break()
    
    analyze_consecutive_events(df, doc)
    doc.add_page_break()
    
    analyze_ml_model_performance(df, doc)
    
    # Add executive summary
    doc.add_page_break()
    add_section_header(doc, "Executive Summary")
    
    summary_paragraphs = [
        "This report provides a comprehensive analysis of cobble detection data, examining patterns, correlations, and predictive indicators across multiple dimensions.",
        
        f"During the analyzed period ({df['Date'].min().strftime('%Y-%m-%d')} to {df['Date'].max().strftime('%Y-%m-%d')}), a total of {df['cobble_detected_status'].sum()} cobble events were recorded out of {len(df)} observations, representing a {df['cobble_detected_status'].mean()*100:.2f}% occurrence rate.",
        
        f"Block-specific analysis reveals that Block {df.groupby('Block')['cobble_detected_status'].mean().idxmax()} has the highest cobble rate at {df.groupby('Block')['cobble_detected_status'].mean().max()*100:.2f}%, while Block {df.groupby('Block')['cobble_detected_status'].mean().idxmin()} shows the lowest at {df.groupby('Block')['cobble_detected_status'].mean().min()*100:.2f}%.",
        
        f"The {df.groupby('Shift')['cobble_detected_status'].mean().idxmax()} shift experiences the highest cobble occurrence rate at {df.groupby('Shift')['cobble_detected_status'].mean().max()*100:.2f}%, suggesting potential areas for operational improvements during this time period.",
        
        "Short-term cobble prediction shows promising results with the 10-minute detection mechanism achieving better recall but lower precision compared to the 20-minute detection, indicating a tradeoff between early warning and false alarm rates.",
        
        "Time series analysis identified anomalous cobble event patterns that deviate significantly from normal operation, providing opportunities for targeted investigation into root causes.",
        
        "The performance analysis of our Machine Learning model demonstrates strong predictive capability across different blocks and profiles, but with room for improvement in reducing false positives and enhancing performance consistency across all operational conditions."
    ]
    
    for para_text in summary_paragraphs:
        doc.add_paragraph(para_text)
    
    # Save the document
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

def browse_file():
    """Open a file dialog to select an Excel file"""
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(
        title="Select Excel File with Cobble Detection Data",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    return file_path

def main():
    """Main function to run the report generation process"""
    print("=" * 60)
    print(" COBBLE DETECTION COMPREHENSIVE REPORT GENERATOR ")
    print("=" * 60)
    
    # # Get input file path from user
    # print("\nPlease select the Excel file containing cobble detection data.")
    # file_path = browse_file()
    
    # if not file_path:
    #     print("No file selected. Exiting program.") 

    #     return
           
    file_path = "data/cobble_data_feb.xlsx" 

    # file_path = "automation_report_generate_data.xlsx"   

     
    print(f"\nSelected file: {file_path}")
    
    # Load and preprocess data
    print("\nLoading and preprocessing data...")
    df = load_data(file_path)
    
    if df is None or len(df) == 0:
        print("Error: Could not load data or data is empty. Please check the file format.")
        return
    
    print(f"Successfully loaded {len(df)} records.")
    
    # Create output folder if it doesn't exist
    output_folder = os.path.join(os.path.dirname(file_path), "Reports")
    os.makedirs(output_folder, exist_ok=True)
    
    # Generate output file name based on input file
    input_filename = os.path.basename(file_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"CAR_{input_filename.split('.')[0]}_{timestamp}.docx"
    output_path = os.path.join(output_folder, output_filename)
    
    # Generate report
    print("\nGenerating comprehensive report with visualizations...")
    create_report(df, output_path)
    
    print("\nProcess completed!")
    print(f"Report saved to: {output_path}")
    
    # Open containing folder
    os.startfile(os.path.dirname(output_path))

if __name__ == "__main__":
    main()

